import os
import docx

def main():
    template_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "srs_template-ieee.docx"))
    print("Loading template from:", template_path)
    
    if not os.path.exists(template_path):
        print("Error: Template file does not exist!")
        return
        
    doc = docx.Document(template_path)
    
    # Locate the target paragraph to insert before
    target_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "Appendix C: To Be Determined List" in p.text:
            target_idx = i
            break
            
    if target_idx == -1:
        print("Error: Could not find 'Appendix C: To Be Determined List' paragraph!")
        return
        
    print(f"Found insertion target at index {target_idx}: '{doc.paragraphs[target_idx].text}'")
    target_p = doc.paragraphs[target_idx]
    
    # 1. Heading 2: B.3 User Interface Mockup and Component Specifications
    h2_1 = target_p.insert_paragraph_before("B.3 User Interface Mockup & UI Component Specifications", style="Heading 2")
    
    # 2. Paragraph: Screen Name
    p_screen = target_p.insert_paragraph_before("Mockup Screen Name: {{ page_name }}", style="Body Text")
    
    # 3. Paragraph: Mockup Image placeholder
    p_img = target_p.insert_paragraph_before("{{ mockup_image }}", style="Normal")
    p_img.alignment = 1 # Center alignment
    
    # 4. Heading 2: B.4 UI Component Specification Table
    h2_2 = target_p.insert_paragraph_before("B.4 UI Component Specification Table", style="Heading 2")
    
    # 5. Table insertion
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    # Fill headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Label'
    hdr_cells[3].text = 'Description'
    
    # Set bold font for headers
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                
    # Fill loop row
    row_cells = table.rows[1].cells
    row_cells[0].text = "{%tr for el in ui_elements %}{{ el.id }}"
    row_cells[1].text = "{{ el.type }}"
    row_cells[2].text = "{{ el.label }}"
    row_cells[3].text = "{{ el.description }}{%tr endfor %}"
    
    # Move the table before target_p
    target_p._element.addprevious(table._element)
    
    # 6. Heading 2: B.5 Detected User Flows
    h2_3 = target_p.insert_paragraph_before("B.5 Detected User Flows", style="Heading 2")
    
    # 7. List item for user flows loop
    p_flow_start = target_p.insert_paragraph_before("{% for flow in detected_user_flows %}", style="Body Text")
    p_flow_item = target_p.insert_paragraph_before("- {{ flow }}", style="List Paragraph")
    p_flow_end = target_p.insert_paragraph_before("{% endfor %}", style="Body Text")
    
    # Insert some empty spacing paragraphs for formatting
    target_p.insert_paragraph_before("")
    
    temp_path = template_path + ".tmp"
    try:
        doc.save(temp_path)
        if os.path.exists(template_path):
            os.remove(template_path)
        os.rename(temp_path, template_path)
        print("Template successfully updated and saved at:", template_path)
    except Exception as e:
        print("Failed to save and replace template:", e)
        if os.path.exists(temp_path):
            os.remove(temp_path)

if __name__ == "__main__":
    main()
