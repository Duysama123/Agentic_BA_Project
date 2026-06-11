import os
import docx

def main():
    template_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "srs_template-ieee.docx"))
    print("Loading template from:", template_path)
    if not os.path.exists(template_path):
        print("Error: Template file does not exist!")
        return

    doc = docx.Document(template_path)
    
    # Check Table 2 (index 2)
    if len(doc.tables) <= 2:
        print("Error: Table 2 does not exist in template!")
        return
        
    table = doc.tables[2]
    print("Found table 2: rows=", len(table.rows), "cols=", len(table.columns))
    
    # Row 1 cells
    cells = table.rows[1].cells
    print("Current cells text:")
    for i, cell in enumerate(cells):
        print(f"  Cell {i}: {cell.text}")
        
    # Modify cell texts to use {%tr syntax
    cells[0].text = "{%tr for el in ui_elements %}{{ el.id }}"
    cells[1].text = "{{ el.type }}"
    cells[2].text = "{{ el.label }}"
    cells[3].text = "{{ el.description }}{%tr endfor %}"
    
    print("Updating cells text to use %tr syntax...")
    for i, cell in enumerate(cells):
        print(f"  Updated Cell {i}: {cell.text}")
        
    # Save the updated template
    temp_path = template_path + ".tmp"
    try:
        doc.save(temp_path)
        if os.path.exists(template_path):
            os.remove(template_path)
        os.rename(temp_path, template_path)
        print("Template successfully fixed and saved at:", template_path)
    except Exception as e:
        print("Failed to save template:", e)
        if os.path.exists(temp_path):
            os.remove(temp_path)

if __name__ == "__main__":
    main()
