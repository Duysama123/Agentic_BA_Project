import os
import docx

def main():
    template_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "srs_template-ieee.docx"))
    print("Loading template from:", template_path)
    if not os.path.exists(template_path):
        print("Error: Template file does not exist!")
        return

    doc = docx.Document(template_path)
    
    # Check Table 2
    if len(doc.tables) <= 2:
        print("Error: Table 2 does not exist in template!")
        return
        
    old_table = doc.tables[2]
    print("Found old table 2: rows=", len(old_table.rows), "cols=", len(old_table.columns))
    
    # Create new 4-row table
    new_table = doc.add_table(rows=4, cols=4)
    new_table.style = 'Table Grid'
    
    # Row 0: Headers
    hdr_cells = new_table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Label'
    hdr_cells[3].text = 'Description'
    
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                
    # Row 1: Loop start
    new_table.rows[1].cells[0].text = "{%tr for el in ui_elements %}"
    
    # Row 2: Data row
    row_cells = new_table.rows[2].cells
    row_cells[0].text = "{{ el.id }}"
    row_cells[1].text = "{{ el.type }}"
    row_cells[2].text = "{{ el.label }}"
    row_cells[3].text = "{{ el.description }}"
    
    # Row 3: Loop end
    new_table.rows[3].cells[0].text = "{%tr endfor %}"
    
    # Replace old_table with new_table in the XML tree
    old_table._element.addnext(new_table._element)
    old_table._element.getparent().remove(old_table._element)
    
    # Save the updated template
    temp_path = template_path + ".tmp"
    try:
        doc.save(temp_path)
        if os.path.exists(template_path):
            os.remove(template_path)
        os.rename(temp_path, template_path)
        print("Template successfully fixed with 4-row table and saved at:", template_path)
    except Exception as e:
        print("Failed to save template:", e)
        if os.path.exists(temp_path):
            os.remove(temp_path)

if __name__ == "__main__":
    main()
