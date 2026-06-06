import docx
import re
try:
    doc = docx.Document('../srs_template-ieee.docx')
    for p in doc.paragraphs:
        if '{' in p.text or '%' in p.text:
            print("Paragraph:", p.text)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if '{' in p.text or '%' in p.text:
                        print("Cell:", p.text)
except Exception as e:
    print(e)
