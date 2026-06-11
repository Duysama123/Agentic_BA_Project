import os
import json
from docxtpl import DocxTemplate

def generate_srs_docx(ba_data_json: str, template_path: str, output_path: str, diagram_data_json: str = None, vision_data_json: str = None, image_bytes: bytes = None) -> str:
    """
    Load data from BA Agent into the IEEE DOCX template.
    If diagram_data_json is provided, appends to the DOCX.
    If vision_data_json/image_bytes are provided, integrates them.
    Returns the path to the generated DOCX.
    """
    data = json.loads(ba_data_json)
    
    from datetime import datetime
    if 'date_created' not in data:
        data['date_created'] = datetime.now().strftime("%B %d, %Y")
    if 'author' not in data:
        data['author'] = "Agentic BA System"
    if 'organization' not in data:
        data['organization'] = "E-Commerce Project Team"
    
    # Initialize the template
    doc = DocxTemplate(template_path)
    
    # Check if the user has added diagram variables to the template
    template_vars = doc.get_undeclared_template_variables()
    
    append_at_end = True
    append_vision_at_end = True
    temp_files = []
    
    # Process diagrams
    if diagram_data_json:
        diag_data = json.loads(diagram_data_json)
        data['diagram_explanation'] = diag_data.get("diagram_explanation", "")
        
        if 'flowchart_image' in template_vars or 'sequence_image' in template_vars:
            append_at_end = False
            import requests
            import zlib
            import base64
            import tempfile
            from docxtpl import InlineImage
            from docx.shared import Inches
            
            def fetch_diagram_image(mermaid_text: str) -> bytes:
                mermaid_text = mermaid_text.strip()
                if mermaid_text.startswith("```mermaid"): mermaid_text = mermaid_text[10:]
                elif mermaid_text.startswith("```"): mermaid_text = mermaid_text[3:]
                if mermaid_text.endswith("```"): mermaid_text = mermaid_text[:-3]
                mermaid_text = mermaid_text.strip()
                
                # 1. Try mermaid.ink
                try:
                    encoded = base64.urlsafe_b64encode(mermaid_text.encode('utf-8')).decode('ascii')
                    r = requests.get(f"https://mermaid.ink/img/{encoded}", timeout=5)
                    if r.status_code == 200:
                        return r.content
                except Exception as e:
                    print("mermaid.ink fetch failed, trying kroki...", e)
                    
                # 2. Try kroki.io as fallback
                try:
                    compressed = zlib.compress(mermaid_text.encode('utf-8'), 9)
                    encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
                    r = requests.get(f"https://kroki.io/mermaid/png/{encoded}", timeout=5)
                    if r.status_code == 200:
                        return r.content
                except Exception as e:
                    print("Kroki fetch failed too:", e)
                    
                return None

            fc_code = diag_data.get("flowchart_diagram", "")
            if fc_code and 'flowchart_image' in template_vars:
                img_data = fetch_diagram_image(fc_code)
                if img_data:
                    try:
                        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                        tmp.write(img_data)
                        tmp.close()
                        temp_files.append(tmp.name)
                        data['flowchart_image'] = InlineImage(doc, tmp.name, width=Inches(6.0))
                    except Exception as e:
                        print("Failed to embed flowchart image:", e)

            seq_code = diag_data.get("sequence_diagram", "")
            if seq_code and 'sequence_image' in template_vars:
                img_data = fetch_diagram_image(seq_code)
                if img_data:
                    try:
                        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                        tmp.write(img_data)
                        tmp.close()
                        temp_files.append(tmp.name)
                        data['sequence_image'] = InlineImage(doc, tmp.name, width=Inches(6.0))
                    except Exception as e:
                        print("Failed to embed sequence image:", e)

    # Process vision data (mockup and components)
    if vision_data_json:
        try:
            v_data = json.loads(vision_data_json)
            data['page_name'] = v_data.get('page_name', 'UI')
            
            raw_elements = v_data.get('elements', [])
            ui_elements = []
            for el in raw_elements:
                ui_elements.append({
                    'id': el.get('id', '') if isinstance(el, dict) else getattr(el, 'id', ''),
                    'type': el.get('type', '') if isinstance(el, dict) else getattr(el, 'type', ''),
                    'label': el.get('label', '') if isinstance(el, dict) else getattr(el, 'label', ''),
                    'description': el.get('description', '') if isinstance(el, dict) else getattr(el, 'description', '')
                })
            data['ui_elements'] = ui_elements
            data['detected_user_flows'] = v_data.get('detected_user_flows', [])
            
            if 'mockup_image' in template_vars or 'ui_elements' in template_vars:
                append_vision_at_end = False
                if image_bytes:
                    import tempfile
                    from docxtpl import InlineImage
                    from docx.shared import Inches
                    try:
                        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                        tmp.write(image_bytes)
                        tmp.close()
                        temp_files.append(tmp.name)
                        data['mockup_image'] = InlineImage(doc, tmp.name, width=Inches(6.0))
                    except Exception as ie:
                        print("Failed to save mockup image for inline template:", ie)
        except Exception as ve:
            print("Failed to parse vision data context:", ve)

    print("DEBUG - data keys:", data.keys())
    
    # Render the document with the context
    doc.render(data)
    
    # Save the output
    doc.save(output_path)
    
    for tmp_file in temp_files:
        try:
            os.unlink(tmp_file)
        except:
            pass
            
    # Append extras if available and not included in template
    if (diagram_data_json and append_at_end) or (vision_data_json and append_vision_at_end):
        append_extras_to_docx(
            output_path, 
            diagram_data_json if append_at_end else None, 
            vision_data_json if append_vision_at_end else None, 
            image_bytes if append_vision_at_end else None
        )
        
    return output_path

def append_extras_to_docx(docx_path: str, diagram_data_json: str = None, vision_data_json: str = None, image_bytes: bytes = None):
    import base64
    import zlib
    import requests
    import tempfile
    from docx import Document
    from docx.shared import Inches
    import json
    import os

    try:
        doc = Document(docx_path)
        
        # 1. Append diagrams if provided and not handled inline
        if diagram_data_json:
            diag_data = json.loads(diagram_data_json)
            doc.add_page_break()
            doc.add_heading('System Diagrams (Auto-Generated)', level=1)

            def fetch_diagram_image(mermaid_text: str) -> bytes:
                mermaid_text = mermaid_text.strip()
                if mermaid_text.startswith("```mermaid"): mermaid_text = mermaid_text[10:]
                elif mermaid_text.startswith("```"): mermaid_text = mermaid_text[3:]
                if mermaid_text.endswith("```"): mermaid_text = mermaid_text[:-3]
                mermaid_text = mermaid_text.strip()
                
                # 1. Try mermaid.ink
                try:
                    encoded = base64.urlsafe_b64encode(mermaid_text.encode('utf-8')).decode('ascii')
                    r = requests.get(f"https://mermaid.ink/img/{encoded}", timeout=5)
                    if r.status_code == 200:
                        return r.content
                except Exception as e:
                    print("mermaid.ink fetch failed, trying kroki...", e)
                    
                # 2. Try kroki.io as fallback
                try:
                    compressed = zlib.compress(mermaid_text.encode('utf-8'), 9)
                    encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
                    r = requests.get(f"https://kroki.io/mermaid/png/{encoded}", timeout=5)
                    if r.status_code == 200:
                        return r.content
                except Exception as e:
                    print("Kroki fetch failed too:", e)
                    
                return None

            # Fetch and add Flowchart
            fc_code = diag_data.get("flowchart_diagram", "")
            if fc_code:
                img_data = fetch_diagram_image(fc_code)
                if img_data:
                    try:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                            tmp.write(img_data)
                            tmp_path = tmp.name
                        doc.add_heading('Flowchart Diagram', level=2)
                        doc.add_picture(tmp_path, width=Inches(6.0))
                        os.unlink(tmp_path)
                    except Exception as e:
                        print("Failed to add flowchart image:", e)

            # Fetch and add Sequence
            seq_code = diag_data.get("sequence_diagram", "")
            if seq_code:
                img_data = fetch_diagram_image(seq_code)
                if img_data:
                    try:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                            tmp.write(img_data)
                            tmp_path = tmp.name
                        doc.add_heading('Sequence Diagram', level=2)
                        doc.add_picture(tmp_path, width=Inches(6.0))
                        os.unlink(tmp_path)
                    except Exception as e:
                        print("Failed to add sequence image:", e)

            explanation = diag_data.get("diagram_explanation", "")
            if explanation:
                doc.add_heading('Diagram Explanation', level=2)
                doc.add_paragraph(explanation)

        # 2. Append vision if provided and not handled inline
        if vision_data_json:
            v_data = json.loads(vision_data_json)
            doc.add_page_break()
            doc.add_heading('Appendix C: User Interface Mockup & Component Specifications', level=1)
            
            page_name = v_data.get('page_name', 'UI Screen')
            doc.add_heading(f'Mockup Screen: {page_name}', level=2)
            
            if image_bytes:
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                        tmp.write(image_bytes)
                        tmp_path = tmp.name
                    doc.add_picture(tmp_path, width=Inches(6.0))
                    os.unlink(tmp_path)
                except Exception as e:
                    print("Failed to append mockup picture:", e)
                    
            # Add table
            doc.add_heading('UI Component Specifications', level=2)
            elements = v_data.get('elements', [])
            if elements:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'ID'
                hdr_cells[1].text = 'Type'
                hdr_cells[2].text = 'Label'
                hdr_cells[3].text = 'Description'
                
                # Make header bold
                for cell in hdr_cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.bold = True
                
                for el in elements:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(el.get('id', ''))
                    row_cells[1].text = str(el.get('type', ''))
                    row_cells[2].text = str(el.get('label', '') or '')
                    row_cells[3].text = str(el.get('description', ''))
                    
            # Add user flows
            flows = v_data.get('detected_user_flows', [])
            if flows:
                doc.add_heading('Detected User Flows', level=2)
                for flow in flows:
                    doc.add_paragraph(flow, style='List Paragraph')

        doc.save(docx_path)
    except Exception as e:
        print("Failed to append extras:", e)

def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """
    Convert DOCX to PDF.
    Tries LibreOffice first (for Linux/Streamlit Cloud).
    Falls back to docx2pdf (for Windows/Mac with MS Word installed).
    """
    import platform
    import subprocess
    import os

    # 1. Try LibreOffice on Linux (Streamlit Cloud)
    if platform.system() == 'Linux':
        try:
            outdir = os.path.dirname(pdf_path)
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", docx_path, "--outdir", outdir],
                check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE
            )
            # libreoffice saves the file as same name but .pdf
            expected_pdf = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
            if expected_pdf != pdf_path and os.path.exists(expected_pdf):
                os.rename(expected_pdf, pdf_path)
            return True
        except Exception as e:
            print(f"LibreOffice conversion failed: {e}")

    # 2. Try docx2pdf for local Windows/Mac users
    try:
        from docx2pdf import convert
        import pythoncom
        
        # docx2pdf uses COM on Windows; if called from a background thread (like Streamlit sometimes does),
        # COM needs to be initialized.
        pythoncom.CoInitialize()
        
        convert(docx_path, pdf_path)
        return True
    except Exception as e:
        print(f"Error converting to PDF (MS Word might not be installed or accessible): {e}")
        return False
